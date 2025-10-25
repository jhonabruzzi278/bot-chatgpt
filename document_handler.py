#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Manejador de documentos para el bot.
Soporta PDF, Word, Excel y archivos de texto.
"""

import logging
import io
import csv
from typing import Optional, Tuple
from pathlib import Path
from config.settings import setup_rotating_logger

# Importaciones opcionales (se verifican al usar)
try:
    import PyPDF2  # type: ignore
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import docx  # type: ignore
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl  # type: ignore
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

logger = setup_rotating_logger("document-handler", "document-handler.log")

class DocumentHandler:
    """Procesa diferentes tipos de documentos."""
    
    def __init__(self):
        """Inicializa el manejador de documentos."""
        self.max_file_size = 20 * 1024 * 1024  # 20 MB
        self.supported_extensions = {
            'pdf': self._process_pdf,
            'txt': self._process_text,
            'docx': self._process_word,
            'doc': self._process_word,
            'xlsx': self._process_excel,
            'xls': self._process_excel,
            'csv': self._process_csv
        }
    
    def is_supported(self, filename: str) -> bool:
        """Verifica si el archivo es soportado."""
        extension = Path(filename).suffix.lower().lstrip('.')
        return extension in self.supported_extensions
    
    def get_supported_formats(self) -> str:
        """Retorna lista de formatos soportados."""
        return ", ".join([f".{ext}" for ext in self.supported_extensions.keys()])
    
    async def process_document(self, file_bytes: bytes, filename: str) -> Tuple[bool, str, Optional[str]]:
        """
        Procesa un documento y extrae su contenido.
        
        Returns:
            Tuple[bool, str, Optional[str]]: (éxito, mensaje, contenido_extraído)
        """
        try:
            # Verificar tamaño
            if len(file_bytes) > self.max_file_size:
                size_mb = len(file_bytes) / (1024 * 1024)
                return False, f"❌ Archivo muy grande: {size_mb:.1f} MB. Máximo: 20 MB", None
            
            # Obtener extensión
            extension = Path(filename).suffix.lower().lstrip('.')
            
            if extension not in self.supported_extensions:
                return False, f"❌ Formato no soportado. Formatos válidos: {self.get_supported_formats()}", None
            
            # Procesar según tipo
            processor = self.supported_extensions[extension]
            content = await processor(file_bytes, filename)
            
            if not content or len(content.strip()) == 0:
                return False, "❌ No se pudo extraer texto del documento", None
            
            # Limitar longitud del contenido
            max_chars = 30000  # ~8000 tokens aproximadamente
            if len(content) > max_chars:
                content = content[:max_chars] + "\n\n[... Documento truncado por longitud ...]"
            
            return True, f"✅ Documento procesado exitosamente ({len(content)} caracteres)", content
            
        except Exception as e:
            logger.error(f"Error procesando documento {filename}: {e}")
            return False, f"❌ Error procesando documento: {str(e)}", None
    
    async def _process_pdf(self, file_bytes: bytes, filename: str) -> str:
        """Extrae texto de un PDF."""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 no está instalado. Instala con: pip install PyPDF2")
        
        try:
            pdf_file = io.BytesIO(file_bytes)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_content = []
            num_pages = len(pdf_reader.pages)
            
            for page_num, page in enumerate(pdf_reader.pages, 1):
                text = page.extract_text()
                if text.strip():
                    text_content.append(f"--- Página {page_num}/{num_pages} ---\n{text}\n")
            
            full_text = "\n".join(text_content)
            
            if not full_text.strip():
                raise ValueError("No se pudo extraer texto del PDF")
            
            return full_text
            
        except ImportError:
            raise ImportError("PyPDF2 no está instalado. Instala con: pip install PyPDF2")
        except Exception as e:
            logger.error(f"Error procesando PDF {filename}: {e}")
            raise
    
    async def _process_text(self, file_bytes: bytes, filename: str) -> str:
        """Procesa archivos de texto plano."""
        try:
            # Intentar diferentes codificaciones
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    return file_bytes.decode(encoding)
                except UnicodeDecodeError:
                    continue
            
            # Si ninguna funciona, usar utf-8 con reemplazo de errores
            return file_bytes.decode('utf-8', errors='replace')
            
        except Exception as e:
            logger.error(f"Error procesando texto {filename}: {e}")
            raise
    
    async def _process_word(self, file_bytes: bytes, filename: str) -> str:
        """Extrae texto de documentos Word."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx no está instalado. Instala con: pip install python-docx")
        
        try:
            doc_file = io.BytesIO(file_bytes)
            doc = docx.Document(doc_file)
            
            text_content = []
            
            # Extraer párrafos
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
            
            # Extraer tablas
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    if row_text.strip():
                        text_content.append(row_text)
            
            full_text = "\n\n".join(text_content)
            
            if not full_text.strip():
                raise ValueError("No se pudo extraer texto del documento Word")
            
            return full_text
            
        except ImportError:
            raise ImportError("python-docx no está instalado. Instala con: pip install python-docx")
        except Exception as e:
            logger.error(f"Error procesando Word {filename}: {e}")
            raise
    
    async def _process_excel(self, file_bytes: bytes, filename: str) -> str:
        """Extrae datos de archivos Excel."""
        if not EXCEL_AVAILABLE:
            raise ImportError("openpyxl no está instalado. Instala con: pip install openpyxl")
        
        try:
            excel_file = io.BytesIO(file_bytes)
            workbook = openpyxl.load_workbook(excel_file, data_only=True)
            
            text_content = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_content.append(f"\n=== Hoja: {sheet_name} ===\n")
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                    if row_text.strip():
                        text_content.append(row_text)
            
            full_text = "\n".join(text_content)
            
            if not full_text.strip():
                raise ValueError("No se pudo extraer datos del archivo Excel")
            
            return full_text
            
        except ImportError:
            raise ImportError("openpyxl no está instalado. Instala con: pip install openpyxl")
        except Exception as e:
            logger.error(f"Error procesando Excel {filename}: {e}")
            raise
    
    async def _process_csv(self, file_bytes: bytes, filename: str) -> str:
        """Procesa archivos CSV."""
        try:
            # Decodificar contenido
            text_content = await self._process_text(file_bytes, filename)
            
            # Parsear CSV
            csv_reader = csv.reader(io.StringIO(text_content))
            
            rows = []
            for row in csv_reader:
                rows.append(" | ".join(row))
            
            return "\n".join(rows)
            
        except Exception as e:
            logger.error(f"Error procesando CSV {filename}: {e}")
            raise

# Instancia global
document_handler = DocumentHandler()
